# app.py
from flask import Flask, render_template, request, redirect, url_for, send_file, flash
from flask_sqlalchemy import SQLAlchemy
from io import BytesIO
from datetime import datetime
import os
from werkzeug.utils import secure_filename

# ReportLab para PDF
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image

# Word
from docx import Document
from docx.shared import Inches

# Excel (.xlsx)
import openpyxl
import pandas as pd

# Matplotlib para generar el medidor
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib import patches
import math

app = Flask(__name__)
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///audit.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.secret_key = 'cambiar-esta-clave'

UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

db = SQLAlchemy(app)

# Estados permitidos
STATES = [
    "Fortalezas",
    "Hallazgos",
    "No conformidad menor",
    "No conformidad mayor",
    "Observaciones"
]

# MODELOS
class Empresa(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    nombre = db.Column(db.String(200), nullable=False, unique=True)
    # relationship: empresa.questions

class Question(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    text = db.Column(db.Text, nullable=False)
    state = db.Column(db.String(100), nullable=True)
    observation = db.Column(db.Text, nullable=True)
    empresa_id = db.Column(db.Integer, db.ForeignKey('empresa.id'), nullable=True)
    empresa = db.relationship("Empresa", backref="questions")

class AuditReport(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    empresa_id = db.Column(db.Integer, db.ForeignKey('empresa.id'), nullable=True)
    empresa = db.relationship("Empresa")
    empresa_nombre = db.Column(db.String(200), nullable=True)
    auditor_nombre = db.Column(db.String(200), nullable=True)
    firma_auditor = db.Column(db.String(200), nullable=True)
    firma_empresa = db.Column(db.String(200), nullable=True)
    auditor_text = db.Column(db.Text, nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

with app.app_context():
    db.create_all()

# ---------- Helpers ----------
def compute_compliance(summary_counts):
    """
    summary_counts: dict with keys = STATES and values = counts
    Returns: (pct_compliance_float, status_string)
    Logic:
      cumplimiento = (Fortalezas + Observaciones) / total_answered * 100
      thresholds:
        >=85% -> 'Cumple'
        60-84.99 -> 'Cumplimiento Parcial'
        <60 -> 'No Cumple'
    If total_answered == 0, returns (0.0, 'Sin respuestas').
    """
    favorable = summary_counts.get("Fortalezas", 0) + summary_counts.get("Observaciones", 0)
    total = sum(summary_counts.values())
    if total == 0:
        return 0.0, "Sin respuestas"
    pct = (favorable / total) * 100.0
    pct_rounded = round(pct, 1)
    if pct >= 85.0:
        status = "Cumple"
    elif pct >= 60.0:
        status = "Cumplimiento Parcial"
    else:
        status = "No Cumple"
    return pct_rounded, status

def draw_gauge(pct, filename, title=None, size=(8,4)):
    """
    Dibuja un medidor semicircular y lo guarda como PNG en `filename`.
    pct: 0..100
    filename: ruta de salida
    """
    # Normalize
    pct = max(0.0, min(100.0, float(pct)))

    fig = plt.figure(figsize=size, dpi=150)
    ax = fig.add_subplot(111, polar=False)
    ax.set_xlim(-1.1, 1.1)
    ax.set_ylim(-0.05, 1.2)
    ax.axis('off')

    # Draw colored wedges from red (0%) to green (100%)
    # We'll draw 5 segments corresponding to ranges (0-20,20-40,40-60,60-80,80-100)
    ranges = [(0,20),(20,40),(40,60),(60,80),(80,100)]
    colors_list = ['#e53935','#fb8c00','#fdd835','#c6e48b','#2e8b57']  # red->orange->yellow->lightgreen->green

    # Draw segments as Wedge patches centered at (0,0) but we will convert to polar positions by computing arcs
    # We'll place wedges by using patches.Wedge with angles in degrees.
    inner_r = 0.25
    outer_r = 1.0
    start_angle = 180  # left
    for i, r in enumerate(ranges):
        seg_start = start_angle - (r[0] / 100.0) * 180.0
        seg_end = start_angle - (r[1] / 100.0) * 180.0
        wedge = patches.Wedge((0,0), outer_r, theta1=seg_end, theta2=seg_start, width=outer_r - inner_r, facecolor=colors_list[i], edgecolor='white')
        ax.add_patch(wedge)

    # Draw ticks and labels (0,25,50,75,100)
    for val in [0,25,50,75,100]:
        angle_deg = 180 - (val/100.0) * 180.0
        angle_rad = math.radians(angle_deg)
        x_outer = math.cos(angle_rad) * (outer_r + 0.02)
        y_outer = math.sin(angle_rad) * (outer_r + 0.02)
        x_inner = math.cos(angle_rad) * (inner_r - 0.02)
        y_inner = math.sin(angle_rad) * (inner_r - 0.02)
        ax.plot([x_inner, x_outer], [y_inner, y_outer], lw=0, color='black')
        # labels
        lx = math.cos(angle_rad) * (outer_r + 0.12)
        ly = math.sin(angle_rad) * (outer_r + 0.12)
        ax.text(lx, ly, f"{val}", horizontalalignment='center', verticalalignment='center', fontsize=8)

    # Draw needle
    # Map pct 0->180deg, 100->0deg
    angle_deg = 180 - (pct/100.0) * 180.0
    angle_rad = math.radians(angle_deg)
    needle_len = (outer_r + inner_r)/2 + 0.05
    nx = math.cos(angle_rad) * needle_len
    ny = math.sin(angle_rad) * needle_len
    ax.plot([0, nx], [0, ny], lw=4, color='#0b3d91')  # needle
    # center circle
    center_circle = patches.Circle((0,0), 0.06, facecolor='#0b3d91', zorder=5)
    ax.add_patch(center_circle)

    # Title and percent text
    if title:
        ax.text(0, 1.05, title, horizontalalignment='center', fontsize=12, weight='bold')
    #ax.text(0, -0.02, f"{pct:.1f}% - {('Sin respuestas' if pct==0 and pct!=100 else '')}", horizontalalignment='center', fontsize=10)

    plt.savefig(filename, bbox_inches='tight', pad_inches=0.1)
    plt.close(fig)
    return filename

# ---------------- RUTAS GENERALES ----------------
@app.route('/')
def index():
    empresas = Empresa.query.order_by(Empresa.nombre).all()
    questions = Question.query.order_by(Question.id).all()
    return render_template('index.html', questions=questions, states=STATES, empresas=empresas)

# ---------------- CRUD PREGUNTAS ----------------
@app.route('/question/new', methods=['GET','POST'])
def new_question():
    empresas = Empresa.query.order_by(Empresa.nombre).all()
    if request.method == 'POST':
        text = request.form.get('text','').strip()
        empresa_id = request.form.get('empresa_id')
        if not text:
            flash('La pregunta no puede estar vacía', 'danger')
            return redirect(url_for('new_question'))
        q = Question(text=text)
        if empresa_id:
            try:
                q.empresa_id = int(empresa_id)
            except:
                q.empresa_id = None
        db.session.add(q)
        db.session.commit()
        flash('Pregunta creada', 'success')
        if q.empresa_id:
            return redirect(url_for('empresa_questions', empresa_id=q.empresa_id))
        return redirect(url_for('index'))
    return render_template('question_form.html', empresas=empresas)

@app.route('/question/<int:q_id>/edit', methods=['GET','POST'])
def edit_question(q_id):
    q = Question.query.get_or_404(q_id)
    if request.method == 'POST':
        if 'text' in request.form and request.form.get('text') is not None:
            q.text = request.form.get('text','').strip() or q.text
            db.session.commit()
            flash('Pregunta actualizada', 'success')
            return redirect(url_for('empresa_questions', empresa_id=q.empresa_id) if q.empresa_id else url_for('index'))
        q.state = request.form.get('state')
        q.observation = request.form.get('observation','').strip()
        db.session.commit()
        flash('Respuesta guardada', 'success')
        return redirect(url_for('empresa_questions', empresa_id=q.empresa_id) if q.empresa_id else url_for('index'))
    return render_template('question_edit.html', q=q, states=STATES)

@app.route('/question/<int:q_id>/delete', methods=['POST'])
def delete_question(q_id):
    q = Question.query.get_or_404(q_id)
    empresa_id = q.empresa_id
    db.session.delete(q)
    db.session.commit()
    flash('Pregunta eliminada', 'success')
    if empresa_id:
        return redirect(url_for('empresa_questions', empresa_id=empresa_id))
    return redirect(url_for('index'))

# ---------------- CRUD EMPRESAS ----------------
@app.route('/empresas')
def empresas_list():
    empresas = Empresa.query.order_by(Empresa.nombre).all()
    return render_template('empresas.html', empresas=empresas)

@app.route('/empresa/new', methods=['GET','POST'])
def empresa_new():
    if request.method == 'POST':
        nombre = request.form.get('nombre','').strip()
        if not nombre:
            flash('El nombre es obligatorio', 'danger')
            return redirect(url_for('empresa_new'))
        if Empresa.query.filter_by(nombre=nombre).first():
            flash('Ya existe una empresa con ese nombre', 'danger')
            return redirect(url_for('empresa_new'))
        e = Empresa(nombre=nombre)
        db.session.add(e)
        db.session.commit()
        flash('Empresa creada', 'success')
        return redirect(url_for('empresas_list'))
    return render_template('empresa_form.html', empresa=None)

@app.route('/empresa/<int:empresa_id>/edit', methods=['GET','POST'])
def empresa_edit(empresa_id):
    e = Empresa.query.get_or_404(empresa_id)
    if request.method == 'POST':
        nombre = request.form.get('nombre','').strip()
        if not nombre:
            flash('El nombre es obligatorio', 'danger')
            return redirect(url_for('empresa_edit', empresa_id=empresa_id))
        e.nombre = nombre
        db.session.commit()
        flash('Empresa actualizada', 'success')
        return redirect(url_for('empresas_list'))
    return render_template('empresa_form.html', empresa=e)

@app.route('/empresa/<int:empresa_id>/delete', methods=['POST'])
def empresa_delete(empresa_id):
    e = Empresa.query.get_or_404(empresa_id)
    db.session.delete(e)
    db.session.commit()
    flash('Empresa eliminada', 'success')
    return redirect(url_for('empresas_list'))

# ---------------- SUBIR PREGUNTAS (Excel .xlsx) ----------------
@app.route('/empresa/<int:empresa_id>/upload', methods=['GET','POST'])
def empresa_upload(empresa_id):
    e = Empresa.query.get_or_404(empresa_id)
    if request.method == 'POST':
        file = request.files.get('file')
        if not file or file.filename == '':
            flash('Selecciona un archivo .xlsx', 'danger')
            return redirect(url_for('empresa_upload', empresa_id=empresa_id))
        filename = secure_filename(f"{datetime.utcnow().timestamp()}_{file.filename}")
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        try:
            wb = openpyxl.load_workbook(filepath, read_only=True)
            ws = wb.active
            added = 0
            for row in ws.iter_rows(min_row=1, values_only=True):
                if not row:
                    continue
                cell = row[0]
                if cell is None:
                    continue
                text = str(cell).strip()
                if text == '':
                    continue
                q = Question(text=text, empresa_id=e.id)
                db.session.add(q)
                added += 1
            db.session.commit()
            flash(f'Se importaron {added} preguntas para la empresa "{e.nombre}"', 'success')
        except Exception as ex:
            flash(f'Error al leer el archivo: {ex}', 'danger')
        # volver a la vista de preguntas de la empresa (no a empresas)
        return redirect(url_for('empresa_questions', empresa_id=empresa_id))
    return render_template('upload_questions.html', empresa=e)

# ---------------- PÁGINAS POR EMPRESA (diligenciar, informe, export) ----------------
@app.route('/empresa/<int:empresa_id>/preguntas')
def empresa_questions(empresa_id):
    e = Empresa.query.get_or_404(empresa_id)
    questions = Question.query.filter_by(empresa_id=empresa_id).order_by(Question.id).all()
    return render_template('empresa_questions.html', empresa=e, questions=questions, states=STATES)

@app.route('/empresa/<int:empresa_id>/diligenciar', methods=['GET','POST'])
def empresa_diligenciar(empresa_id):
    e = Empresa.query.get_or_404(empresa_id)
    questions = Question.query.filter_by(empresa_id=empresa_id).order_by(Question.id).all()
    if request.method == 'POST':
        for q in questions:
            q.state = request.form.get(f'state_{q.id}')
            q.observation = request.form.get(f'obs_{q.id}','').strip()
        db.session.commit()
        flash('Respuestas guardadas', 'success')
        # volver a preguntas de la empresa (no al informe)
        return redirect(url_for('empresa_questions', empresa_id=empresa_id))
    return render_template('empresa_diligenciar.html', empresa=e, questions=questions, states=STATES)

@app.route('/empresa/<int:empresa_id>/audit', methods=['GET','POST'])
def empresa_audit(empresa_id):
    e = Empresa.query.get_or_404(empresa_id)
    questions = Question.query.filter_by(empresa_id=empresa_id).order_by(Question.id).all()
    summary = {s: 0 for s in STATES}
    for q in questions:
        if q.state in summary:
            summary[q.state] += 1
    if request.method == 'POST':
        auditor_nombre = request.form.get('auditor_nombre','').strip()
        auditor_text = request.form.get('auditor_text','').strip()
        # firmas opcionales
        firma_auditor_file = request.files.get('firma_auditor')
        firma_empresa_file = request.files.get('firma_empresa')
        firma_auditor_path = None
        firma_empresa_path = None

        if firma_auditor_file and firma_auditor_file.filename != '':
            filename = secure_filename(f"{datetime.utcnow().timestamp()}_auditor_{firma_auditor_file.filename}")
            firma_auditor_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            firma_auditor_file.save(firma_auditor_path)

        if firma_empresa_file and firma_empresa_file.filename != '':
            filename = secure_filename(f"{datetime.utcnow().timestamp()}_empresa_{firma_empresa_file.filename}")
            firma_empresa_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            firma_empresa_file.save(firma_empresa_path)

        ar = AuditReport(
            empresa_id=e.id,
            empresa_nombre=e.nombre,
            auditor_nombre=auditor_nombre,
            auditor_text=auditor_text,
            firma_auditor=firma_auditor_path,
            firma_empresa=firma_empresa_path
        )
        db.session.add(ar)
        db.session.commit()

        # elegir si generar word o pdf
        if request.form.get('generate') == 'word':
            return redirect(url_for('export_word_by_empresa', empresa_id=empresa_id, report_id=ar.id))
        return redirect(url_for('export_pdf_by_empresa', empresa_id=empresa_id, report_id=ar.id))
    return render_template('empresa_audit_form.html', empresa=e, questions=questions, summary=summary)

# ---------------- EXPORTAR WORD / PDF POR EMPRESA ----------------
@app.route('/empresa/<int:empresa_id>/report/word/<int:report_id>')
def export_word_by_empresa(empresa_id, report_id):
    ar = AuditReport.query.get_or_404(report_id)
    questions = Question.query.filter_by(empresa_id=empresa_id).order_by(Question.id).all()

    summary = {s: 0 for s in STATES}
    for q in questions:
        if q.state in summary:
            summary[q.state] += 1

    # calcular cumplimiento
    pct, status = compute_compliance(summary)

    # generar medidor (gauge) y guardarlo en uploads
    gauge_filename = os.path.join(app.config['UPLOAD_FOLDER'], f"gauge_emp{empresa_id}_r{report_id}_{int(datetime.utcnow().timestamp())}.png")
    try:
        draw_gauge(pct, gauge_filename, title=f"Cumplimiento: {pct}%")
    except Exception as ex:
        gauge_filename = None

    doc = Document()
    doc.add_heading('INFORME DE AUDITORÍA', level=1)
    doc.add_paragraph(f"Fecha: {ar.created_at.strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Empresa: {ar.empresa_nombre or '(no definida)'}")
    doc.add_paragraph(f"Auditor: {ar.auditor_nombre or '(no definido)'}")
    doc.add_paragraph("")

    doc.add_heading("Resumen de hallazgos", level=2)
    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = "Estado"
    hdr[1].text = "Cantidad"
    for s in STATES:
        row = table.add_row().cells
        row[0].text = s
        row[1].text = str(summary[s])

    # Indicador de cumplimiento (Word) - texto y medidor si se generó
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run(f"Cumplimiento: {pct}%  —  Resultado: {status}")
    run.bold = True
    doc.add_paragraph("")

    if gauge_filename and os.path.exists(gauge_filename):
        try:
            doc.add_picture(gauge_filename, width=Inches(6))
        except Exception:
            # si falla, no interrumpe
            pass

    doc.add_paragraph("")
    doc.add_heading("Resultados Detallados", level=2)
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "#"
    hdr[1].text = "Pregunta"
    hdr[2].text = "Estado"
    hdr[3].text = "Observación"
    for q in questions:
        row = table.add_row().cells
        row[0].text = str(q.id)
        row[1].text = q.text
        row[2].text = q.state or "(sin estado)"
        row[3].text = q.observation or "(sin observación)"

    doc.add_paragraph("")
    doc.add_heading("Informe Final del Auditor", level=2)
    doc.add_paragraph(ar.auditor_text or "(sin observaciones)")

    doc.add_paragraph("")
    doc.add_heading("Firmas", level=2)
    signatures = doc.add_table(rows=2, cols=2)
    sig_titles = signatures.rows[0].cells
    sig_titles[0].text = "Auditor"
    sig_titles[1].text = "Empresa"
    sig_images = signatures.rows[1].cells
    if ar.firma_auditor:
        try:
            sig_images[0].paragraphs[0].add_run().add_picture(ar.firma_auditor, width=Inches(2))
        except Exception:
            sig_images[0].text = "(Imagen no disponible)"
    else:
        sig_images[0].text = "______________________"
    if ar.firma_empresa:
        try:
            sig_images[1].paragraphs[0].add_run().add_picture(ar.firma_empresa, width=Inches(2))
        except Exception:
            sig_images[1].text = "(Imagen no disponible)"
    else:
        sig_images[1].text = "______________________"

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    filename = f"informe_{ar.empresa_nombre or 'empresa'}_{ar.created_at.strftime('%Y%m%d_%H%M%S')}.docx"
    return send_file(buffer, as_attachment=True, download_name=filename, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@app.route('/empresa/<int:empresa_id>/report/pdf/<int:report_id>')
def export_pdf_by_empresa(empresa_id, report_id):
    ar = AuditReport.query.get_or_404(report_id)
    questions = Question.query.filter_by(empresa_id=empresa_id).order_by(Question.id).all()

    summary = {s: 0 for s in STATES}
    for q in questions:
        if q.state in summary:
            summary[q.state] += 1

    # calcular cumplimiento
    pct, status = compute_compliance(summary)

    # generar medidor (gauge) y guardarlo en uploads
    gauge_filename = os.path.join(app.config['UPLOAD_FOLDER'], f"gauge_emp{empresa_id}_r{report_id}_{int(datetime.utcnow().timestamp())}.png")
    try:
        draw_gauge(pct, gauge_filename, title=f"Cumplimiento: {pct}%")
    except Exception:
        gauge_filename = None

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter,
                            rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(name="Title", parent=styles["Title"], alignment=1, fontSize=16, leading=20, spaceAfter=12)
    heading_style = ParagraphStyle(name="Heading", parent=styles["Heading2"], fontSize=12, leading=14, spaceAfter=6)
    small_style = ParagraphStyle(name="Small", parent=styles["BodyText"], fontSize=9, leading=11)
    cell_style = ParagraphStyle(name="Cell", parent=styles["BodyText"], fontSize=9, leading=11, spaceBefore=2, spaceAfter=2)
    auditor_style = ParagraphStyle(name="AuditorText", parent=styles["BodyText"], fontSize=11, leading=16, spaceBefore=6, spaceAfter=6)
    indicator_style = ParagraphStyle(name="Indicator", parent=styles["BodyText"], fontSize=11, leading=14, spaceBefore=6, spaceAfter=6)

    body = []
    body.append(Paragraph("INFORME DE AUDITORÍA", title_style))
    body.append(Paragraph(f"Fecha: {ar.created_at.strftime('%Y-%m-%d %H:%M:%S')}", styles["Normal"]))
    body.append(Spacer(1, 12))
    body.append(Paragraph(f"Empresa: {ar.empresa_nombre or '(no definida)'}", styles["Normal"]))
    body.append(Paragraph(f"Auditor: {ar.auditor_nombre or '(no definido)'}", styles["Normal"]))
    body.append(Spacer(1, 12))

    # Resumen de hallazgos
    body.append(Paragraph("Resumen de hallazgos", heading_style))
    body.append(Spacer(1, 12))
    summary_data = [[Paragraph("<b>Estado</b>", small_style), Paragraph("<b>Cantidad</b>", small_style)]]
    for s in STATES:
        summary_data.append([Paragraph(s, small_style), Paragraph(str(summary[s]), small_style)])
    summary_table = Table(summary_data, colWidths=[180, 80], hAlign='LEFT')
    summary_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
        ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
    ]))
    body.append(summary_table)
    body.append(Spacer(1, 12))

    # Resultados detallados
    body.append(Paragraph("Resultados Detallados", heading_style))
    body.append(Spacer(1, 12))
    table_data = [
        [Paragraph("<b>#</b>", small_style),
         Paragraph("<b>Pregunta</b>", small_style),
         Paragraph("<b>Estado</b>", small_style),
         Paragraph("<b>Observación</b>", small_style)]
    ]
    for q in questions:
        table_data.append([
            Paragraph(str(q.id), small_style),
            Paragraph(q.text or "(sin texto)", cell_style),
            Paragraph(q.state or "(sin estado)", cell_style),
            Paragraph(q.observation or "(sin observación)", cell_style)
        ])
    details_table = Table(table_data, colWidths=[30, 260, 90, 160], repeatRows=1, hAlign='LEFT')
    details_table.setStyle(TableStyle([
        ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
        ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
    ]))
    body.append(details_table)
    body.append(Spacer(1, 18))

    # Indicador de cumplimiento (PDF) - texto
    indicator_text = f"Cumplimiento de la auditoria: {pct}%  —  Resultado obtenido: {status}"
    body.append(Paragraph(indicator_text, indicator_style))
    body.append(Spacer(1, 8))

    # Insertar imagen del gauge si existe
    if gauge_filename and os.path.exists(gauge_filename):
        try:
            # Insert at a reasonable display size
            img = Image(gauge_filename, width=400, height=200)
            body.append(img)
            body.append(Spacer(1, 12))
        except Exception:
            pass


    # Informe final
    body.append(Paragraph("Informe Final del Auditor", heading_style))
    body.append(Spacer(1, 12))
    auditor_text = (ar.auditor_text or "(sin observaciones)").replace("\n", "<br/>")
    body.append(Paragraph(auditor_text, auditor_style))
    body.append(Spacer(1, 18))

    # Firmas
    firma_row = []
    if ar.firma_auditor:
        try:
            firma_row.append(Image(ar.firma_auditor, width=200, height=50))
        except Exception:
            firma_row.append(Paragraph(" (firma auditor no disponible) ", styles["Normal"]))
    else:
        firma_row.append(Paragraph("____________________", styles["Normal"]))

    if ar.firma_empresa:
        try:
            firma_row.append(Image(ar.firma_empresa, width=200, height=50))
        except Exception:
            firma_row.append(Paragraph(" (firma empresa no disponible) ", styles["Normal"]))
    else:
        firma_row.append(Paragraph("____________________", styles["Normal"]))

    firma_table = Table([firma_row], colWidths=[270, 270], hAlign='CENTER')
    body.append(firma_table)

    nombre_row = [
        Paragraph("Auditor", styles["Normal"]),
        Paragraph("Empresa", styles["Normal"])
    ]
    nombre_table = Table([nombre_row], colWidths=[270, 270], hAlign='CENTER')
    body.append(nombre_table)

    doc.build(body)
    buffer.seek(0)
    filename = f"informe_{ar.empresa_nombre or 'empresa'}_{ar.created_at.strftime('%Y%m%d_%H%M%S')}.pdf"
    return send_file(buffer, as_attachment=True, download_name=filename, mimetype='application/pdf')

# ---------------- ADICIONES: agregar manual, eliminar todas, exportar a XLSX ----------------
@app.route("/empresa/<int:empresa_id>/add_manual", methods=["POST"])
def add_question_manual(empresa_id):
    text = request.form.get("question_text", "").strip()
    if text:
        q = Question(text=text, empresa_id=empresa_id)
        db.session.add(q)
        db.session.commit()
        flash("Pregunta agregada", "success")
    else:
        flash("La pregunta no puede estar vacía", "danger")
    return redirect(url_for("empresa_questions", empresa_id=empresa_id))

@app.route("/empresa/<int:empresa_id>/delete_all", methods=["POST"])
def delete_all_questions(empresa_id):
    Question.query.filter_by(empresa_id=empresa_id).delete()
    db.session.commit()
    flash("Todas las preguntas de la empresa fueron eliminadas", "success")
    return redirect(url_for("empresa_questions", empresa_id=empresa_id))

@app.route("/empresa/<int:empresa_id>/export_excel")
def export_questions_excel(empresa_id):
    questions = Question.query.filter_by(empresa_id=empresa_id).order_by(Question.id).all()
    data = []
    for q in questions:
        data.append({
            "ID": q.id,
            "Pregunta": q.text,
            "Estado": q.state or "",
            "Observación": q.observation or ""
        })
    df = pd.DataFrame(data)

    output = BytesIO()
    if df.empty:
        df = pd.DataFrame(columns=["ID", "Pregunta", "Estado", "Observación"])

    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Preguntas")
    output.seek(0)

    filename = f"Preguntas_{empresa_id}_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.xlsx"
    return send_file(output,
                     as_attachment=True,
                     download_name=filename,
                     mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

# ------------------------------------------------
if __name__ == '__main__':
    app.run(debug=True)
